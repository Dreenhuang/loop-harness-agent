# -*- coding: utf-8 -*-
"""
生成《Loop Agent 项目教学.docx》
面向1年经验产品经理的通俗教学,深度讲解 + 适当延伸
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== 样式辅助 ==============
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_quote(doc, text, label="💡 划重点"):
    """高亮引用块"""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'FFF4CE')
    cell.text = ''
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(f"{label}\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xC0, 0x50, 0x4D)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.size = Pt(11)
    p2.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()  # 空行

def add_code_block(doc, code, lang=""):
    """代码块"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F2F2F2')
    cell.text = ''
    p = cell.paragraphs[0]
    if lang:
        r0 = p.add_run(f"[{lang}]\n")
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    r.font.size = Pt(9)
    doc.add_paragraph()

def add_table_styled(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1F4E79')
    # body
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_bg(cell, 'F2F7FC')
    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("─" * 30)
    r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

# ============== 创建文档 ==============
doc = Document()

# 全局样式:中文宋体,英文Calibri
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ============== 封面 ==============
cover_p = doc.add_paragraph()
cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_p.add_run("\n\n\n\n")
r.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("从 0 到 1 读懂 Loop Agent")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("—— 一个 AI 自动化开发流水线项目的通俗教学课")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
r.italic = True

doc.add_paragraph()
target = doc.add_paragraph()
target.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = target.add_run("🎯 适合人群:1 年经验的产品经理")
r.font.size = Pt(13)

target2 = doc.add_paragraph()
target2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = target2.add_run("📚 阅读时长:约 40 分钟")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()
doc.add_paragraph()
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = intro.add_run("通过剖析一个真实可运行的 AI 自动化开发项目,\n带你看懂 AI 时代产品经理必须懂的 7 大核心概念。")
r.font.size = Pt(12)
r.italic = True
r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_page_break()

# ============== 课程导览 ==============
add_h1(doc, "📖 课程导览:这堂课你会学到什么")
add_para(doc, "这堂课不会教你写代码,也不会教你怎么用 AI 工具。")
add_para(doc, "我们会一起拆解一个真实存在的项目 ——「Loop Agent」(全称 Loop-Harness-Agent,中文名:循环代理)。")
add_quote(doc, "这个项目是干嘛的?一句话:它是一个「AI 时代的产品开发流水线」,把一个产品从需求到上线的整个过程,拆成 16 个角色、10 个阶段、4 道门禁,让 AI 像一支虚拟团队一样帮你把产品做出来。")

add_para(doc, "作为产品经理,看懂这个项目你会获得 3 个核心能力:")
add_bullet(doc, "【能力 1】理解「AI 自动化开发」到底是什么、能做什么、不能做什么")
add_bullet(doc, "【能力 2】看懂复杂项目的「架构图」,能参与讨论技术方案")
add_bullet(doc, "【能力 3】理解 AI Agent 的「协作方式」,将来和工程师、设计师沟通不卡壳")

add_h2(doc, "课程目录(7 节课)")
add_numbered(doc, "第 1 课:这个项目到底是什么?(大白话版)")
add_numbered(doc, "第 2 课:核心概念 1 ——「角色」(Agent Team)")
add_numbered(doc, "第 3 课:核心概念 2 ——「流水线」(Phase + Gate)")
add_numbered(doc, "第 4 课:核心概念 3 ——「黑板」(Blackboard)")
add_numbered(doc, "第 5 课:核心概念 4 ——「MCP Server」(把规则外挂出来)")
add_numbered(doc, "第 6 课:产品经理的「第一性原理」—— 这个项目想解决啥问题?")
add_numbered(doc, "第 7 课:实战启示 —— 我们能从这个项目里学到什么?")

add_divider(doc)

# ============== 第 1 课 ==============
add_h1(doc, "第 1 课 | 这个项目到底是什么?(大白话版)")
add_para(doc, "在拆解前,先说个最朴素的类比,后面所有的概念都建立在这个类比上。")
add_quote(doc, "把「Loop Agent」想成一家虚拟的「AI 数字公司」。\n这家公司有员工、有部门、有流程、有 KPI、还有安检门。\n你(产品经理)只要提交一份「产品需求文档」(PRD),这家公司就能自己把产品做出来交给你。")

add_h2(doc, "1.1 这家「AI 数字公司」的 3 个核心要素")
add_table_styled(doc, ["现实公司里的东西", "Loop Agent 里的对应", "大白话解释"],
[
    ["员工", "16 个 Agent(角色)", "每个 Agent 都是一个「会干某件事的 AI」"],
    ["工作流程", "10 个 Phase(阶段)", "把工作拆成 10 步,一步一步来"],
    ["绩效考核 / 验收", "4 道 Gate(门禁)", "每完成一批,都要过 4 道安检"],
    ["公司公告板", "Blackboard(黑板)", "记录谁做到哪了、有啥问题"],
    ["公司预算", "Budget(预算)", "花钱 / 跑多少步有上限,不能无限烧钱"],
])

add_h2(doc, "1.2 一个最简化的产品开发流程")
add_para(doc, "想象你让这家公司做一个「在线点奶茶小程序」,流程是这样的:")
add_numbered(doc, "Phase 0(初始化):公司先把会议室、规章、员工名单准备好")
add_numbered(doc, "Phase 1(需求):产品经理 + 需求分析师,写出清晰的需求文档")
add_numbered(doc, "Phase 2(交互设计):UX 研究员,画出用户操作流程")
add_numbered(doc, "Phase 3(视觉设计):UI 设计师,出高保真设计稿")
add_numbered(doc, "Phase 4(架构):架构师,定技术方案、数据库、API")
add_numbered(doc, "Phase 5(开发):后端 + 前端工程师,并行写代码")
add_numbered(doc, "Phase 6(质量门禁):3 道安检 —— 代码审查、性能压测、功能测试")
add_numbered(doc, "Phase 7(知识沉淀):把这次踩的坑写下来,下次别再犯")
add_numbered(doc, "Phase 8(文档):把代码、API、用户手册整理成册")
add_numbered(doc, "Phase 9(终审):总审核员签字,确认可以上线")
add_numbered(doc, "Phase 10(部署):运维同学把产品推到线上服务器")

add_quote(doc, "记住这 10 个阶段!你以后看到任何 AI 自动化开发的文章,只要提到「pipeline」「workflow」,基本都是这个套路 —— 拆阶段、设门禁、配角色。")

add_divider(doc)

# ============== 第 2 课 ==============
add_h1(doc, "第 2 课 | 核心概念 1 ——「角色」(Agent Team)")
add_para(doc, "这家「AI 数字公司」一共雇了 16 个员工,按职责分成 4 个层级。")
add_quote(doc, "为什么是 16 个,不是 3 个、不是 50 个?\n因为 16 个刚好覆盖一个产品从「想法」到「上线」所需的全部专业能力,再多就内卷了,再少就有人要干两份活,容易出错。")

add_h2(doc, "2.1 16 个角色的 4 大层级")
add_table_styled(doc, ["层级", "角色", "大白话"],
[
    ["调度层(1 个)", "@Orchestrator(总指挥)", "项目经理,只管调度,不管具体干活"],
    ["决策层(1 个)", "@Product-Manager(产品经理)", "你!负责拍板需求优先级"],
    ["业务层(5 个)", "@Requirements / @UX / @UI / @Architect", "把需求翻译成图纸"],
    ["技术层(5 个)", "@Backend / @Frontend / @Bug-Fixer / @DevOps", "写代码、修 Bug、部署"],
    ["质量层(3 个)", "@Code-Reviewer / @Performance / @Tester", "3 道安检员"],
    ["交付层(2 个)", "@Documenter / @Final-Reviewer", "写文档、做最终签字"],
    ["知识层(1 个)", "@Knowledge-Curator", "把经验沉淀成知识库"],
])

add_h2(doc, "2.2 产品经理视角:看这 16 个角色的 2 个启发")
add_quote(doc, "【启发 1】一个产品需要的「专业岗位」其实是有清单的。\n下次你招聘时,可以参考这 16 个角色,问自己:我的团队缺哪个专业岗位?\n不要什么都让一个「全栈」来扛 —— 现实里 16 个专家,虚拟世界里 16 个 AI,效率都比「全栈」高。")

add_quote(doc, "【启发 2】每个角色都有「准入技能」(bound_skills)。\n一个 @Code-Reviewer 必须掌握 gate1-code-review(代码审查技能)才能上岗。\n这对 PM 的启示:招人 / 找 AI 工具时,要明确「这个岗位的核心技能清单」,而不是「会某样东西就行」。")

add_h2(doc, "2.3 延伸:为什么「总指挥」只调度不干活?")
add_para(doc, "注意看第一行:@Orchestrator(总指挥)只负责「调度」,它不写需求、不写代码、不做设计。")
add_para(doc, "这是为什么?")
add_quote(doc, "如果总指挥也干活,就会出现两个问题:\n① 总指挥累死了,流程就停了;\n② 总指挥自己写的东西,自己审,容易「自欺欺人」,质量门禁形同虚设。\n\n这个原则叫「Maker-Checker 分离」—— 干活的人和检查的人不能是同一个人。\n这是 Loop Agent 里非常核心的一条纪律,也是任何团队管理的通则。")

add_divider(doc)

# ============== 第 3 课 ==============
add_h1(doc, "第 3 课 | 核心概念 2 ——「流水线」(Phase + Gate)")
add_para(doc, "光有 16 个员工不够,你还得规定好「先干啥后干啥」、「啥时候能通过」,否则大家一拥而上,乱套了。")
add_para(doc, "这个项目用「Phase(阶段)+ Gate(门禁)」两个机制搞定。")

add_h2(doc, "3.1 Phase(阶段):把工作切成 10 块")
add_para(doc, "10 个阶段不是随便切的,它们遵循一个规律:从抽象到具体,从想法到产物。")
add_table_styled(doc, ["阶段类别", "包含哪些 Phase", "产物"],
[
    ["想法层", "Phase 0~1(初始化+需求)", "PRD 文档"],
    ["图纸层", "Phase 2~3(交互+视觉)", "设计稿、组件库"],
    ["蓝图层", "Phase 4(架构)", "技术选型、API 规范"],
    ["实施层", "Phase 5(开发)", "源代码"],
    ["验证层", "Phase 6(质量门禁)", "测试报告"],
    ["沉淀层", "Phase 7~8(知识+文档)", "知识库、用户手册"],
    ["发布层", "Phase 9~10(终审+部署)", "上线产品"],
])

add_quote(doc, "产品经理的启发:在你自己规划产品 Roadmap 时,也可以参考这种「想法→图纸→蓝图→实施→验证→沉淀→发布」的 7 段式节奏。不要一上来就写代码,也不要上线了才补文档。")

add_h2(doc, "3.2 Gate(门禁):4 道质量关卡")
add_para(doc, "代码写完了不能直接上线,要过 4 道关卡:")
add_table_styled(doc, ["门禁", "负责角色", "通过标准", "没通过会怎样"],
[
    ["Gate 1:代码审查", "@Code-Reviewer", "0 Blocker + 0 Major 级别问题", "回到 Phase 5 重写"],
    ["Gate 2:性能压测", "@Performance-Engineer", "P95 响应 ≤ 300ms,错误率 ≤ 0.1%", "回到 Phase 5 优化"],
    ["Gate 3:功能测试", "@全栈测试员", "P0/P1 级别 Bug = 0,P2 ≤ 3 个", "回到 Phase 5 修复"],
    ["Gate 4:终审", "@Final-Reviewer", "风险等级 ≤ LOW", "回到对应 Phase 返工"],
])

add_quote(doc, "产品经理的启发:为什么需要 4 道门禁?因为「质量不是测出来的,是流程里长出来的」。\n单点检查容易漏、容易偷懒;4 道关卡,每道只盯一个维度,就能保证既快又稳。\n这同样适用于你做产品 —— 需求评审、设计评审、开发联调、灰度发布,本质都是「门禁」。")

add_h2(doc, "3.3 延伸:为什么门禁不通过要「强制回到 Phase 5」?")
add_para(doc, "看项目规则里的一条铁律:")
add_quote(doc, "「门禁不通过 → 强制回到 Phase 5 修复,绝不允许绕过」\n\n这叫「强制回收」纪律。它的意义是:\n  ① 不让小问题滚成大问题;\n  ② 不让「差不多就行」的心态蔓延;\n  ③ 保证下游环节(部署、终审)永远只接收「合格品」。\n\n对应到产品:你设计的任何一个功能,如果开发完成时没达到你定的「验收标准」,就应该打回去,而不是「先上线再说」。「先上线再说」= 跳过门禁,= 流程崩坏。")

add_divider(doc)

# ============== 第 4 课 ==============
add_h1(doc, "第 4 课 | 核心概念 3 ——「黑板」(Blackboard)")
add_para(doc, "16 个角色、10 个阶段、4 道门禁,这么多人在干活,信息怎么互通?")
add_para(doc, "答案是:一块「黑板」。")
add_quote(doc, "在 AI 自动化项目里,「黑板」就是一个文件 —— 「项目进度记录.md」。\n所有人干完活都必须在上面写一笔:我做了啥、遇到啥问题、下一步干啥。\n下次有人(包括 AI 自己)忘记上下文,看一眼黑板就懂了。")

add_h2(doc, "4.1 黑板长啥样?(看真实代码)")
add_para(doc, "这是项目里黑板的实际格式,每个角色每次更新都要按这个结构来:")
add_code_block(doc, """## 【2026-06-17｜Loop Agent｜Phase 5 完成】

### 1. ✅ 本轮已完成
- 后端 API 全部实现
- 前端页面全部开发

### 2. ⚠️ 本轮不确定项
- 性能是否达标

### 3. ❌ 遗留待解决问题
- 需要优化数据库查询

### 4. 📋 下一轮工作计划
- 进入 Gate 1 代码审查

### 5. 🔄 黑板更新记录
- 新增节点:CODE-001
- 状态变更:DOC-005: in_progress → completed
- 快照更新:CP-002(进度 35% → 42%)""", lang="Markdown")

add_h2(doc, "4.2 黑板解决了 AI 的「健忘病」")
add_para(doc, "你可能听过这个段子:「AI 一觉醒来,啥都不记得」。这是真的 —— AI 的对话长度有限,聊多了它就忘了前面说过啥。")
add_quote(doc, "Loop Agent 的解法很巧妙:把所有「上下文」从 AI 脑子里搬到硬盘上。\n  • AI 不用记所有事,只记当前在做啥;\n  • 历史信息全部在黑板上;\n  • 下次启动 AI(比如新建会话),先读黑板,就能「续上记忆」。\n\n这个套路叫「Ralph 模式」—— 每轮重置上下文,进度全存外部。")

add_h2(doc, "4.3 产品经理的延伸:你也能用「黑板思维」管理团队")
add_quote(doc, "【类比 1】Slack / 飞书群:把信息外化,而不是装在每个人脑子里。\n【类比 2】周报制度:每周每个角色写「我做了啥、要干啥、卡在哪」,这就是团队的黑板。\n【类比 3】Trello / 看板:把工作状态外化成卡片,避免「老板不知道大家在干嘛」。\n\n核心思想就一条:重要信息不要只存在人脑里,要有「外部记忆」。\n  • 信息透明 → 协作效率高\n  • 信息可追溯 → 出问题能复盘\n  • 信息不丢失 → 人员变动不影响项目")

add_divider(doc)

# ============== 第 5 课 ==============
add_h1(doc, "第 5 课 | 核心概念 4 ——「MCP Server」(把规则外挂出来)")
add_para(doc, "这是本课最技术的一个概念,我会用最通俗的方式讲。")

add_h2(doc, "5.1 MCP 是啥?(30 秒搞懂)")
add_quote(doc, "MCP = Model Context Protocol,翻译过来就是「AI 模型上下文协议」。\n你可以把它想成「AI 世界的 USB 接口标准」—— 任何 AI 都能用同一种方式连接任何外部工具,就像所有 U 盘都能插进所有电脑一样。\n\n在 Loop Agent 里,MCP Server 就是这个「USB 接口」的实现,负责让外部 AI 调用 Loop Agent 的 6 个动作。")

add_h2(doc, "5.2 Loop Agent 暴露了哪 6 个动作?")
add_table_styled(doc, ["动作名", "作用", "PM 大白话"],
[
    ["start_loop", "启动整个流水线", "按一下「开始」按钮"],
    ["get_status", "查询当前进度", "看一眼项目看板"],
    ["abort_loop", "中止整个流水线", "紧急刹车"],
    ["spawn_agent", "派一个角色干活", "给张三派活"],
    ["list_agents", "看看有哪些角色", "查花名册"],
    ["save_blackboard", "写一行到黑板", "更新一下进度公告"],
])

add_h2(doc, "5.3 为什么一定要做成 MCP Server?(最关键的一节)")
add_para(doc, "你可能会问:这 6 个动作直接在代码里调用不就行了吗?为啥非要「包成 MCP」?")
add_para(doc, "原因有 3 层:")
add_quote(doc, "【原因 1:流程不依赖某个具体 AI】\n  • 之前:Loop Agent 的规则全靠 AI 自己记(读规则文件),换 AI 就要重新教。\n  • 现在:规则、状态、黑板全在 MCP Server 里,AI 只负责想,规矩由 Server 守。\n  • 效果:今天用 Claude,明天换 GPT,后天用 Gemini,都不用改 Loop Agent。")

add_quote(doc, "【原因 2:AI 不会「累」和「忘」】\n  • 之前:AI 聊久了会忘事、可能偷懒跳过门禁。\n  • 现在:门禁逻辑在 Server 里,Server 不会忘、不会累、不会偷懒。\n  • 效果:每个 Phase 都强制更新黑板、每道门禁都强制检查,流程不会跑偏。")

add_quote(doc, "【原因 3:可以多人 / 多 AI 并行使用】\n  • 之前:一个 AI 一个项目,资源冲突。\n  • 现在:MCP Server 是独立进程,可以被多个 AI 同时连接。\n  • 效果:你团队里 3 个产品经理,可以让 3 个 AI 同时通过 MCP 调用 Loop Agent,各自做各自的项目。")

add_h2(doc, "5.4 看看真实代码:100 行实现一个 MCP Server")
add_para(doc, "Loop Agent 的 MCP Server 主体就 100 来行,核心是这两段:")
add_code_block(doc, """// 1. 创建 MCP Server
const server = new Server({
  name: "loop-agent-orchestrator",
  version: "1.0.0"
}, { capabilities: { tools: {} } });

// 2. 告诉外面:我能干这 6 个活
server.setRequestHandler(ListToolsRequestSchema, async () => ({
  tools: [
    { name: "start_loop", description: "启动流程" },
    { name: "spawn_agent", description: "派活" },
    // ... 4 个
  ]
}));

// 3. 接电话:有人调用,就执行
server.setRequestHandler(CallToolRequestSchema, async (req) => {
  // 干活去...改 state.json、改黑板...
});""", lang="TypeScript")

add_quote(doc, "是不是很简单?这就是 MCP Server 的精髓:\n  ① 用一个标准格式告诉外面「我能干啥」;\n  ② 接到请求后,执行相应操作(改文件、调函数);\n  ③ 把结果返回给调用方。\n\n任何工具想「被 AI 调用」,都可以包成 MCP Server。这是 AI 时代的「产品 API 化」趋势。")

add_h2(doc, "5.5 产品经理的延伸:看懂 MCP,你就能看懂未来 3 年的 AI 工具")
add_quote(doc, "MCP 是 2024 年底由 Anthropic 提出的标准,现在(2026 年)已经被业界普遍采用。\n理解了 MCP,你就能理解:\n  • 为什么所有 AI 工具都在「接 MCP」?\n    → 因为接了 MCP,所有 AI 都能用他家的工具,流量自动来。\n  • 什么是「Agent 化」的产品?\n    → 凡是能被 AI 自动调用的功能(查天气、下单、查数据库),都是 Agent 化的产品。\n  • 你自己做的产品要不要「Agent 化」?\n    → 如果你的产品有「让 AI 帮你点外卖、订机票」这类需求,就需要提供 MCP 接口。\n  • 简单的判断标准:你的产品有没有「可被一句话描述的明确动作」?\n    → 有,就能做 MCP 接口;没有,就是普通 App,暂时不需要。")

add_divider(doc)

# ============== 第 6 课 ==============
add_h1(doc, "第 6 课 | 产品经理的「第一性原理」—— 这个项目想解决啥问题?")
add_para(doc, "学了 5 节课的「怎么做的」,现在回到最根本的问题:「为啥要做这个」?")

add_h2(doc, "6.1 AI 开发的最大痛点(也是这个项目要解决的问题)")
add_para(doc, "在 2024-2026 年,大家用 AI 写代码时普遍遇到 3 个痛点:")
add_quote(doc, "【痛点 1:流程失控】\n  • 症状:AI 一口气写完所有代码,但没设计、没测试、没文档,直接交差。\n  • 结果:代码能跑,但是烂代码,后期改不动。")

add_quote(doc, "【痛点 2:质量不稳】\n  • 症状:有时候 AI 写得很好,有时候写得很烂,完全看心情。\n  • 结果:每次都得人工 review,效率反而更低。")

add_quote(doc, "【痛点 3:协作混乱】\n  • 症状:不同 AI 写出来的代码风格不一样,模块之间对不上,接口不规范。\n  • 结果:项目越大越乱,最后只能推翻重写。")

add_h2(doc, "6.2 Loop Agent 的解法:把「流程」当产品做")
add_para(doc, "Loop Agent 创始人(也就是你看到的这个项目的作者)想了一个妙招:")
add_quote(doc, "如果 AI 写代码不可控,那就别让 AI「自由发挥」,而是给 AI 一套「约束剧本」:\n  • 谁先谁后(10 个 Phase);\n  • 谁干啥(16 个 Agent);\n  • 啥时候算合格(4 道 Gate);\n  • 中间状态记在哪(Blackboard);\n  • 规矩由谁守(MCP Server)。\n\n这就是「Harness」(约束具)的意思 —— 不是「AI 越自由越强」,而是「给 AI 戴上有规矩的笼子,它才能飞得高」。")

add_h2(doc, "6.3 第一性原理:产品经理的 3 个根本启发")
add_quote(doc, "【启发 1:流程本身就是产品】\n  • 你做的产品,用户体验很重要;\n  • 但你做产品的「流程」,同样是产品 —— 它决定了你交付的质量。\n  • 投资优化流程,就是投资质量。")

add_quote(doc, "【启发 2:自动化不是「替代人」,而是「让机器守规矩」】\n  • 常见的错误认知:AI 自动化 = 让人下岗。\n  • 真正的认知:AI 自动化 = 让机器干那些「不能错、不能懒、不能忘」的事(比如门禁、黑板)。\n  • 让人去做「需要创造力、决策力」的事(比如需求、体验)。")

add_quote(doc, "【启发 3:复杂系统的核心是「边界清晰」】\n  • 16 个角色,每个角色只干一件事;\n  • 10 个阶段,每个阶段只产出一个产物;\n  • 4 道门禁,每道只检查一个维度。\n  • 这就是「单一职责」原则 —— 不要让一个角色 / 阶段 / 门禁做太多事,否则一定乱。\n  • 你设计产品功能时,也是这个道理:每个按钮只做一件事,每个页面只解决一个场景。")

add_divider(doc)

# ============== 第 7 课 ==============
add_h1(doc, "第 7 课 | 实战启示 —— 我们能从这个项目里学到什么?")
add_para(doc, "最后一课,我们把项目里的「招式」翻译成 PM 能用的「工作心法」。")

add_h2(doc, "7.1 5 条 PM 可直接用的心法")
add_table_styled(doc, ["项目里的做法", "PM 对应的工作心法", "举例"],
[
    ["16 个 Agent 角色", "画一张「专业岗位图」", "你的团队需要哪些专业岗位?有没有缺?"],
    ["10 个 Phase 阶段", "把项目拆成 7~10 段", "从「想法」到「上线」拆几个阶段,每个阶段产出啥?"],
    ["4 道 Gate 门禁", "在关键节点设「验收会」", "需求评审、设计评审、灰度评审、发布评审"],
    ["Blackboard 黑板", "把信息外化到「公共区」", "周报、Slack、共享文档、Trello"],
    ["MCP Server 外挂", "把规则「系统化」而不是「人化」", "别让规矩只存在老员工脑子里,要有 SOP"],
])

add_h2(doc, "7.2 3 个「避坑提醒」")
add_quote(doc, "【避坑 1】别一上来就「自动化」,先「流程化」\n  • 错误:老板让你用 AI 自动化整个产品,你直接上 AI。\n  • 正确:先把流程标准化(谁干啥、啥时候算完),再考虑用 AI 自动化其中一两个环节。\n  • 顺序:手工 → SOP → 半自动 → 全自动。")

add_quote(doc, "【避坑 2】别忽视「门禁」,别让「差不多就行」蔓延\n  • 错误:产品快上线了,还有 10 个 P2 Bug,你说「先上再说」。\n  • 正确:回到「门禁」环节,该修就修,该延期就延期。\n  • 原则:「差不多就行」永远是质量崩坏的开始。")

add_quote(doc, "【避坑 3】别让「总指挥」亲自下场干活\n  • 错误:产品经理又写需求、又画原型、又做交互、还跟开发对接口。\n  • 正确:产品经理只负责「调度和决策」,不负责「执行」,否则会累死或者「自欺欺人」。\n  • 原则:Maker-Checker 分离 —— 干活的和检查的不能是同一个人。")

add_h2(doc, "7.3 一张图总结:Loop Agent 给 PM 的启示")
add_quote(doc, "【一图流】\n  • 16 个角色 → 你的团队岗位图\n  • 10 个阶段 → 你的 Roadmap 节奏\n  • 4 道门禁 → 你的关键节点评审\n  • 1 块黑板 → 你的周报 / 共享文档\n  • MCP Server → 你的 SOP 制度化\n\n这 5 件事做好了,你的产品开发效率会提升 50% 以上 —— 不需要 AI 也能做到;加上 AI 之后,就是 5 倍 10 倍的提升。")

add_h2(doc, "7.4 给 PM 的 3 本延伸书单")
add_quote(doc, "【书单 1】《团队协作的五大障碍》\n  • 解决:为什么「Maker-Checker 分离」重要?为什么团队要写周报?\n  • 推荐理由:5 个障碍讲透了所有团队问题。")

add_quote(doc, "【书单 2】《精益创业》\n  • 解决:为什么 Loop Agent 要做 10 个阶段、4 道门禁?为什么不能「先上再说」?\n  • 推荐理由:MVP 思维 + 验证循环,和产品开发节奏一模一样。")

add_quote(doc, "【书单 3】《系统之美》\n  • 解决:为什么 Loop Agent 强调「边界清晰」「单一职责」?\n  • 推荐理由:系统思考的入门书,看完你就能看懂所有复杂系统的设计逻辑。")

add_divider(doc)

# ============== 课后寄语 ==============
add_h1(doc, "🎓 课后寄语")
add_para(doc, "恭喜你!你已经看完了 7 节课。")
add_para(doc, "最后送你一段话:")
add_quote(doc, "AI 时代,产品经理的核心能力不是「会用 AI 工具」,而是「能把复杂问题拆成可被 AI 处理的步骤」。\n\nLoop Agent 给你最大的启发不是「AI 多厉害」,而是「好的流程能让 AI 真的发挥价值」。\n\n记住:你不是在和 AI 竞争,你是在用 AI 放大你作为 PM 的专业能力。\n  • 你懂需求,AI 才能写出对的产品;\n  • 你懂流程,AI 才能跑得稳;\n  • 你懂边界,AI 才能不出错。\n\n愿你成为那个「能定义问题」的产品经理 —— 这是 AI 时代最稀缺的能力。", label="🌟 送给你的话")

add_h2(doc, "📌 附录:本项目关键资源")
add_table_styled(doc, ["资源", "路径", "作用"],
[
    ["项目主仓库", "G:\\ai-gongju\\Loop-agent\\", "项目所有源码"],
    ["核心规则", ".trae/rules/loop-agent.md", "主入口规则"],
    ["16 角色配置", ".trae/agents/*.agent.toml", "每个角色的岗位说明书"],
    ["10 阶段蓝图", "workflows/phases/*.json", "每个阶段的详细任务"],
    ["4 道门禁定义", "workflows/gates/*.json", "每道门禁的验收标准"],
    ["MCP Server 代码", "loop-agent-mcp/src/index.ts", "MCP Server 入口(仅 300 行)"],
    ["状态机核心", "loop-agent-engine/orchestrator.ts", "调度逻辑"],
    ["黑板模板", "blackboard/templates/项目进度记录.md", "黑板格式模板"],
])

add_para(doc, " ")
add_para(doc, "本课程基于 Loop Agent 项目 v1.2 版本编写,日期 2026-06-17。", italic=True, size=9)
add_para(doc, "讲解风格:通俗、生活化、有类比,避免学术堆砌。", italic=True, size=9)

# ============== 保存 ==============
output_path = r"G:\ai-gongju\Loop-agent\Loop Agent 项目教学.docx"
doc.save(output_path)
print(f"✅ 教学文档已生成: {output_path}")
print(f"   文件大小: {os.path.getsize(output_path) / 1024:.1f} KB")
print(f"   章节数: 7 课 + 课程导览 + 课后寄语 + 附录")
print(f"   预计阅读时长: 40 分钟")
